from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from django.shortcuts import render, redirect
from pipelines import pipeline
import yake

import wikipedia
# Create your views here.
from myapp.models import*

def home(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    return render(request, 'admin/aindex.html')

def login(request):
    return render(request, 'Login_index.html')

def login_post(request):
    uname=request.POST['name']
    pswrd=request.POST['password']

    log=Login.objects.filter(Name=uname,Password=pswrd)
    if log.exists():
        log1=Login.objects.get(Name=uname,Password=pswrd)
        request.session['lid']=log1.id
        if log1.Usertype=='admin':
            return HttpResponse('''<script>alert('Logined');window.location='/myapp/home/'</script>''')
        elif log1.Usertype == 'User':
            return redirect('/myapp/user_home/')
        else:
            return HttpResponse('''<script>alert("invalid credentials");window.location="/myapp/login/"</script>''')
    else:
        return HttpResponse('''<script>alert("invalid credentials");window.location="/myapp/login/"</script>''')


def cng_pswrd(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    return render(request, 'admin/change_password.html')

def cng_pswrd_pst(request):
    pswd = request.POST['password']
    npswd = request.POST['n_password']
    cpswd = request.POST['c_password']
    log = Login.objects.filter(Password=pswd)
    if log.exists():
        log1 = Login.objects.get(Password=pswd, id=request.session['lid'])
        if npswd == cpswd:
            log1 = Login.objects.filter(Password=pswd, id=request.session['lid']).update(Password=cpswd)
            return HttpResponse('''<script>alert("successfully changed");window.location="/myapp/login/"</script>''')

        else:
            return HttpResponse(
                '''<script>alert("password incorrect");window.location="/myapp/cng_pswrd/"</script>''')
    else:
        return HttpResponse(
            '''<script>alert("password incorrect");window.location="/myapp/cng_pswrd/"</script>''')

def view_userprofile(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    data = User.objects.all()
    return render(request, 'admin/view-user.html', {'dt': data})

def complaint(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    data = Complaints.objects.all()
    return render(request, 'admin/complaints.html', {'dt':data})


def reply(request, id):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    return render(request, 'admin/replay.html', {'id':id})

def reply_pst(request):
    reply = request.POST['reply']
    id = request.POST['id']

    robj = Complaints.objects.get(id=id)
    robj.replay = reply
    robj.status = "replied"
    robj.save()
    return HttpResponse('''<script>alert("Replyed");window.location='/myapp/complaint/'</script>''')

def review(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    data = Reviews.objects.all()
    return render(request, 'admin/view_review.html', {'dt':data})

def review_pst(request):
    frm = request.POST['from']
    to = request.POST['to']
    data = Reviews.objects.filter(date__range=[frm,to])
    return render(request, 'admin/view_review.html', {'dt':data})

def alogout(request):
    request.session['lid'] = ''
    return render(request,'Login_index.html')

################ USER#################



def user_home(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    return render(request, 'user/index.html')


def register(request):
    return render(request, 'user/sindex.html')

def register_pst(request):
    name = request.POST['name']
    phone = request.POST['phone']
    email = request.POST['email']
    dob = request.POST['dob']
    gender = request.POST['gender']
    place = request.POST['place']
    post = request.POST['post']
    district = request.POST['district']
    state = request.POST['state']
    pin = request.POST['pin']
    photo = request.FILES['photo']
    password = request.POST['password']
    c_password = request.POST['cpassword']

    if password==c_password:

        lobj = Login()
        lobj.Name = email
        lobj.Password = c_password
        lobj.Usertype = "User"
        lobj.save()

        robj = User()
        robj.Name = name
        robj.Phone = phone
        robj.Email = email
        robj.Dob = dob
        robj.Gender = gender
        robj.Place = place
        robj.Post = post
        robj.District = district
        robj.State = state
        robj.Pin = pin

        from datetime import datetime
        date = datetime.now().strftime("%Y-%m-%d %H-%M-%S") + ".jpg"
        fs = FileSystemStorage()
        fs.save(date, photo)
        path = fs.url(date)
        robj.Photo = path

        robj.Photo = path
        robj.Password = password
        robj.C_password = c_password
        robj.LOGIN=lobj
        robj.save()

        return HttpResponse('''<script>alert("Registered Succesfully");window.location="/myapp/login/"</script>''')

    else:
        return HttpResponse('''<script>alert("Mismatched");window.location="/myapp/register/"</script>''')


def u_cng_pswrd(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    return render(request, 'user/u_change_pswrd.html')

def u_cng_pswrd_pst(request):
    pswd = request.POST['password']
    npswd = request.POST['n_password']
    cpswd = request.POST['c_password']
    log = Login.objects.filter(Password=pswd)
    if log.exists():
        log1 = Login.objects.get(Password=pswd, id=request.session['lid'])
        if npswd == cpswd:
            log1 = Login.objects.filter(Password=pswd, id=request.session['lid']).update(Password=cpswd)
            return HttpResponse('''<script>alert("successfully changed");window.location="/myapp/login/"</script>''')

        else:
            return HttpResponse(
                '''<script>alert("password incorrect");window.location="/myapp/u_cng_pswrd/"</script>''')
    else:
        return HttpResponse(
            '''<script>alert("password incorrect");window.location="/myapp/u_cng_pswrd/"</script>''')


def view_profile(request):
    if request.session['lid']=='':
        return redirect('/myapp/login/')
    data= User.objects.get(LOGIN_id=request.session['lid'])
    return render(request,'user/view_profile.html', {'dt':data})

def edit_prof(request):
    data = User.objects.get(LOGIN_id=request.session['lid'])
    return render(request,'user/edt_prof.html',{'dt':data})

def edit_prof_pst(request):
    # photo = request.POST['photo']
    name = request.POST['name']
    phone = request.POST['phone']
    email = request.POST['email']
    dob = request.POST['dob']
    id = request.POST['id']



    eobj = User.objects.get(id=id)
    if 'photo' in request.FILES:
        photo = request.FILES['photo']

        from datetime import datetime
        date = datetime.now().strftime("%Y-%m-%d %H-%M-%S") + ".jpg"
        fs = FileSystemStorage()
        fs.save(date, photo)
        path = fs.url(date)
        eobj.Photo = path
    eobj.Name = name
    eobj.Phone = phone
    eobj.Email = email
    eobj.Dob = dob
    eobj.save()



    return HttpResponse('''<script>alert("Updated");window.location="/myapp/view_profile/"</script>''')








# @app.route('/tutor_alg')
def user_alg(request):
    return render(request,'user/alg.html')

def summ(txt):
    import transformers
    from transformers import pipeline

    to_tokenize=txt
    summarizer = pipeline("summarization")
    summarized = summarizer(to_tokenize, min_length=75, max_length=300)

    print(summarized)

    for i in summarized:
        print(i)

    return summarized


def user_alg_post(request):
    message=request.POST['textarea']
    language = "en"
    max_ngram_size = 1
    deduplication_threshold = 0.9
    numOfKeywords = 5
    custom_kw_extractor = yake.KeywordExtractor(lan=language, n=max_ngram_size,
                                                     dedupLim=deduplication_threshold, top=numOfKeywords,
                                                     features=None)
    keywords = custom_kw_extractor.extract_keywords(message)
    # wikipedia.set_lang("hi")

    wikicontent=[]



    g=message

    for i in keywords:
        print(i[0])
        try:
            from wikipedia import wikipedia
            s = wikipedia.summary(str(i[0]))
            print(s)

            s=str(s).replace("'","")
            s=str(s).replace('"',"")
            s=str(s).replace(';',"")

            # g= g +s
            wikicontent.append(s)
        except:
            wikicontent.append("")
            print("error")

    g = str(g).replace("'", "")
    g = str(g).replace('"', "")
    g = str(g).replace(';', "")
    qa = pipeline("question-generation")
    s = qa(g)
    print('aaa', s)

    import os
    import docx
    d = ""
    m = 1
    for i in s:
        d = d + str(m) + "." + i['question'] + "\n" + "Answer: " + i['answer'] + "\n\n\n\n";
        m = m + 1
    save_path = "C:\\Users\\anusr\\PycharmProjects\\Easy_learn\\media\\QnA\\"
    save_path2 = "C:\\Users\\anusr\\PycharmProjects\\Easy_learn\\media\\recording\\"
    # save_path = "C:\\Users\\dell\\PycharmProjects\\easy_learn\\static\\QnA\\"



    # for each in result:
    import time
    time = str(time.time())
    finalpath = (save_path+time)
    # finalpath = (os.path.join(save_path, time))
    finalpath2 = finalpath + ".docx"
    mydoc = docx.Document()
    mydoc.add_paragraph(d)
    # mydoc.add_paragraph((s.replace('\x00','')))
    mydoc.save(finalpath2)

    sm=summ(g)

    docfile= "/media/QnA/"+time+".docx"

    import gtts
    # from playsound import playsound

    # make a request to google to get synthesis
    t1 = gtts.gTTS(sm[0]['summary_text'])
    from datetime import  datetime
    fname= datetime.now().strftime("%Y%m%d%H%M%S")+".mp3"
     # save the audio file
    # t1.save("C:\\Users\\dell\\PycharmProjects\\easy_learn\\static\\recording\\"+fname)

    date=datetime.now().strftime("%Y%m%d-%H%M%S")+fname
    # fs = FileSystemStorage()
    # fs.save('recording/'+fname, t1)
    t1.save("C:\\Users\\anusr\\PycharmProjects\\Easy_learn\\media\\recording\\"+fname)

    reco="/media/recording/"+fname


    qry = uploads()
    qry.voicefilename =reco
    qry.qafilename =docfile
    from datetime import datetime
    qry.date = datetime.now().strftime("%Y-%m-%d")
    qry.LOGIN_id = request.session['lid']
    qry.save()

    return render(request,"user/alg_result.html",{'data':keywords, 'wikicontent':wikicontent, 'qa':s, 'sm':sm})



# @app.route('/tutor_file_upload_post', methods=['POST'])
def user_file_upload_post(request):
    from datetime import datetime
    file=request.FILES['file']
    date=datetime.now().strftime("%Y%m%d-%H%M%S")+file.name
    fs = FileSystemStorage()
    fs.save('uploaded_file/'+date, file)
    # path = fs.url(date)
    # file.save("C:\\Users\\dell\\PycharmProjects\\easy_learn\\static\\uploaded_file\\"+date)
    path="/media/uploaded_file/"+date

    if '.pdf' in file.name:
        import PyPDF2
        pdfFileObj = open(r"C:\Users\anusr\PycharmProjects\Easy_learn\media\uploaded_file\\"+date, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        pageObj = pdfReader.getPage(0)
        print(pageObj.extractText())
        text=pageObj.extractText()
        pdfFileObj.close()
        return render(request,'user/alg.html',{'data':text})

    elif '.docx' in file.name:
        import docx

        doc = docx.Document(r"C:\Users\anusr\PycharmProjects\Easy_learn\media\uploaded_file\\"+date)
        allText = ""
        for docpara in doc.paragraphs:
            allText= allText+ docpara.text
        print(allText)
        return render(request,'user/alg.html',{'data':allText})
    return HttpResponse('''<script>alert("File format not supported");window.location="/myapp/user_alg/"</script>''')


# @app.route('/tutor_view_files')
def user_view_files(request):
    res=uploads.objects.filter(LOGIN_id=request.session['lid'])
    return render(request,'user/Tutor_View_Files.html',{'data':res})

def sendreviewrating(request):
   return render(request, 'user/sendreviewrating.html')


def sendreviewrating_post(request):
    rev=request.POST['textarea']
    rating=request.POST['rating']
    id=request.POST['aid']
    from datetime import datetime
    robj=Reviews()
    robj.review=rev
    robj.rating=rating
    robj.date=datetime.now().strftime("%Y-%m-%d")
    robj.USER=User.objects.get(LOGIN_id=request.session['lid'])
    robj.save()
    return HttpResponse('''<script>alert("successfully send");window.location="/myapp/user_home/"</script>''')


def logout(request):
    request.session['lid'] = ''
    return render(request,'Login_index.html')